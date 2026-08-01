# -*- coding: utf-8 -*-
from pathlib import Path
from math import pi
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage,
    PageBreak, KeepTogether
)

ROOT = Path(r"C:\Users\Administrator\workspace_ccstheia\test1_2.2_recovered")
OUT = ROOT / "tmp" / "design-report"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "车载平衡滚球运动控制系统_H题_设计报告.docx"
PDF_PATH = OUT / "车载平衡滚球运动控制系统_H题_设计报告.pdf"

SIMSUN = r"C:\Windows\Fonts\simsun.ttc"
SIMHEI = r"C:\Windows\Fonts\simhei.ttf"
MSYH = r"C:\Windows\Fonts\msyh.ttc"
FONT_REG = ImageFont.truetype(SIMSUN, 34)
FONT_SMALL = ImageFont.truetype(SIMSUN, 28)
FONT_BOLD = ImageFont.truetype(SIMHEI, 34)
FONT_TITLE = ImageFont.truetype(SIMHEI, 42)

# ---------- diagram helpers ----------
def rounded(draw, box, text, font=FONT_REG, fill="white", outline="black", radius=18, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    x1,y1,x2,y2=box
    bbox=draw.multiline_textbbox((0,0), text, font=font, spacing=6, align="center")
    tw,th=bbox[2]-bbox[0], bbox[3]-bbox[1]
    draw.multiline_text(((x1+x2-tw)/2,(y1+y2-th)/2), text, font=font, fill="black", spacing=6, align="center")

def arrow(draw, p1, p2, width=4, fill="black"):
    draw.line([p1,p2], fill=fill, width=width)
    import math
    ang=math.atan2(p2[1]-p1[1],p2[0]-p1[0])
    size=15
    pts=[]
    for a in (ang+2.6,ang-2.6):
        pts.append((p2[0]+size*math.cos(a),p2[1]+size*math.sin(a)))
    draw.polygon([p2,pts[0],pts[1]], fill=fill)

def save_system_arch(path):
    im=Image.new("RGB",(1600,720),"white"); d=ImageDraw.Draw(im)
    rounded(d,(610,265,990,455),"MSPM0G3507\n主控制器",FONT_TITLE,fill="#f2f2f2",width=4)
    boxes=[
        ((40,70,390,205),"LineFollower_6CH\nI²C0 循迹反馈"),
        ((40,275,390,410),"K230 视觉模块\nUART1 球位置/速度"),
        ((40,480,390,615),"KEY2/KEY3\n模式启动与急停"),
        ((1210,45,1560,180),"DRV8870×2\n左右轮驱动"),
        ((1210,225,1560,360),"D36A+步进电机\n摆杆角度执行"),
        ((1210,405,1560,540),"OLED / 无线链路\n时间、状态与诊断"),
        ((610,555,990,690),"MS42CG 编码器\n执行器内环反馈"),
    ]
    for b,t in boxes: rounded(d,b,t,FONT_REG,fill="white")
    for y in (137,342,547): arrow(d,(390,y),(610,330 if y<240 else 360 if y<450 else 400))
    for y in (112,292,472): arrow(d,(990,330 if y<200 else 360 if y<390 else 400),(1210,y))
    arrow(d,(800,555),(800,455))
    d.text((530,22),"车载平衡滚球运动控制系统总体结构",font=FONT_TITLE,fill="black")
    im.save(path,dpi=(220,220))

def save_control(path):
    im=Image.new("RGB",(1600,690),"white"); d=ImageDraw.Draw(im)
    d.text((540,20),"循迹与滚球双闭环控制结构",font=FONT_TITLE,fill="black")
    # top loop
    rounded(d,(35,105,245,205),"黑线位置",FONT_REG)
    rounded(d,(315,90,570,220),"六路归一化\n加权误差 eₗ",FONT_REG,fill="#f7f7f7")
    rounded(d,(650,90,900,220),"循迹 PID\n+曲线限速",FONT_REG,fill="#f7f7f7")
    rounded(d,(980,90,1240,220),"左右轮PWM\n+速度PI",FONT_REG,fill="#f7f7f7")
    rounded(d,(1320,105,1560,205),"小车运动",FONT_REG)
    for a,b in [((245,155),(315,155)),((570,155),(650,155)),((900,155),(980,155)),((1240,155),(1320,155))]: arrow(d,a,b)
    arrow(d,(1440,205),(1440,260)); arrow(d,(1440,260),(445,260)); arrow(d,(445,260),(445,220))
    d.text((930,265),"编码器速度反馈",font=FONT_SMALL,fill="black")
    # bottom loop
    rounded(d,(35,405,245,505),"目标球位 xᵣ",FONT_REG)
    rounded(d,(315,380,570,530),"球位置 PD/PID\n外环",FONT_REG,fill="#f7f7f7")
    rounded(d,(650,380,900,530),"期望杆角 θᵈ\n限幅/抗饱和",FONT_REG,fill="#f7f7f7")
    rounded(d,(980,380,1240,530),"D36A步进\n角度内环",FONT_REG,fill="#f7f7f7")
    rounded(d,(1320,405,1560,505),"球—杆系统",FONT_REG)
    for a,b in [((245,455),(315,455)),((570,455),(650,455)),((900,455),(980,455)),((1240,455),(1320,455))]: arrow(d,a,b)
    arrow(d,(1440,505),(1440,610)); arrow(d,(1440,610),(445,610)); arrow(d,(445,610),(445,530))
    d.text((820,615),"K230球位反馈；MS42CG角度反馈",font=FONT_SMALL,fill="black")
    im.save(path,dpi=(220,220))

def save_flow(path):
    im=Image.new("RGB",(1500,820),"white"); d=ImageDraw.Draw(im)
    d.text((470,18),"主程序状态机与安全处理流程",font=FONT_TITLE,fill="black")
    cx=750
    ys=[100,230,370,520,675]
    texts=["上电初始化\n全部执行器保持安全态", "读取按键/传感器/通信\n检查新鲜度与故障", "按模式进入 REQ-002 或 REQ-003\n另一模式强制锁止", "周期控制\n循迹PID/速度PI 或 球位外环/角度内环", "完成：冻结计时并停车\n故障：立即STOP并记录原因"]
    sizes=[(420,120),(520,130),(600,135),(640,135),(580,120)]
    boxes=[]
    for y,t,(w,h) in zip(ys,texts,sizes):
        b=(cx-w//2,y,cx+w//2,y+h); boxes.append(b); rounded(d,b,t,FONT_REG,fill="#f7f7f7" if len(boxes)>1 else "white")
    for i in range(len(boxes)-1): arrow(d,((boxes[i][0]+boxes[i][2])//2,boxes[i][3]),((boxes[i+1][0]+boxes[i+1][2])//2,boxes[i+1][1]))
    rounded(d,(35,300,350,465),"无效条件\nI²C失败/丢线\nK230超时/CRC错\n编码器停滞/超期",FONT_SMALL,fill="white")
    arrow(d,(350,382),(450,382))
    rounded(d,(1150,300,1465,465),"安全动作\nPWM=0\nD36A失能\n清积分/冻结计时\n等待人工确认",FONT_SMALL,fill="white")
    arrow(d,(1050,382),(1150,382))
    im.save(path,dpi=(220,220))

def save_calibration(path):
    white=[1921,1514,1830,1604,1850,1607]
    black=[3038,2797,3242,2400,2899,2336]
    im=Image.new("RGB",(1500,520),"white"); d=ImageDraw.Draw(im)
    d.text((410,15),"六路循迹传感器当前标定表对比",font=FONT_TITLE,fill="black")
    x0,y0=120,420; w=1200; h=310; maxv=3500
    d.line((x0,y0,x0+w,y0),fill="black",width=3); d.line((x0,y0-h,x0,y0),fill="black",width=3)
    for v in range(0,3501,500):
        y=y0-h*v/maxv; d.line((x0-8,y,x0+w,y),fill="#dddddd",width=1); d.text((35,y-15),str(v),font=FONT_SMALL,fill="black")
    group=w/6
    for i,(a,b) in enumerate(zip(white,black)):
        gx=x0+i*group+35
        ha=h*a/maxv; hb=h*b/maxv
        d.rectangle((gx,y0-ha,gx+45,y0),fill="#bdbdbd",outline="black")
        d.rectangle((gx+60,y0-hb,gx+105,y0),fill="#333333",outline="black")
        d.text((gx+25,y0+15),f"CH{i+1}",font=FONT_SMALL,fill="black")
    d.rectangle((1230,95,1275,130),fill="#bdbdbd",outline="black"); d.text((1290,95),"白底",font=FONT_SMALL,fill="black")
    d.rectangle((1230,150,1275,185),fill="#333333",outline="black"); d.text((1290,150),"黑线",font=FONT_SMALL,fill="black")
    im.save(path,dpi=(220,220))

FIG_SYS=OUT/'fig-system.png'; FIG_CTRL=OUT/'fig-control.png'; FIG_FLOW=OUT/'fig-flow.png'; FIG_CAL=OUT/'fig-calibration.png'
save_system_arch(FIG_SYS); save_control(FIG_CTRL); save_flow(FIG_FLOW); save_calibration(FIG_CAL)

# ---------- shared content ----------
TITLE="车载平衡滚球运动控制系统（H题）"
DATE="2026年8月1日"
ABSTRACT=("针对循线小车运动时钢球在凹槽摆杆上的位置稳定问题，设计了一种以MSPM0G3507为核心的分层双闭环控制系统。"
"小车采用六路红外传感器加权循迹、差速驱动与双轮速度校正；K230视觉模块输出钢球位置和速度，D36A驱动步进电机调节摆杆，MS42CG编码器构成执行器角度内环。"
"软件以状态机统一管理启动、计时、停车、失效保护和模式互锁，并采用限幅、抗饱和、数据新鲜度检查及物理断能策略提高可靠性。报告给出了运动学与球杆模型、硬件接口、程序流程及分阶段测试方法，为满足单圈时间、停车精度和±1 cm球位误差指标提供实现依据。")
KEYWORDS="关键词：循迹小车；球杆系统；视觉测量；双闭环控制；MSPM0G3507"

# ---------- DOCX helpers ----------
def set_run_font(run, east='宋体', size=12, bold=None):
    run.font.name=east
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size=Pt(size)
    if bold is not None: run.bold=bold

def shade_cell(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_cell_text(cell, text, size=9.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=align
    p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=Pt(14)
    r=p.add_run(str(text)); set_run_font(r,'宋体',size,bold)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def add_page_number(section):
    footer=section.footer; footer.is_linked_to_previous=False
    p=footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=p.add_run(); set_run_font(run,'宋体',10.5)
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text=' PAGE '
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.extend([fldChar1,instrText,fldChar2])

def add_para(doc,text='',style=None,align=WD_ALIGN_PARAGRAPH.JUSTIFY,first=True,size=12,bold=False,space_after=0,line=22):
    p=doc.add_paragraph(style=style); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(0); pf.space_after=Pt(space_after); pf.line_spacing=Pt(line)
    if first: pf.first_line_indent=Cm(0.74)
    r=p.add_run(text); set_run_font(r,'宋体',size,bold)
    return p

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    pf=p.paragraph_format; pf.space_before=Pt(4 if level==1 else 2); pf.space_after=Pt(1); pf.line_spacing=Pt(18)
    r=p.add_run(text); set_run_font(r,'黑体',14 if level==1 else 12,True)
    return p

def add_equation(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=Pt(18)
    r=p.add_run(text); set_run_font(r,'Times New Roman',11,False); r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    return p

def add_caption(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=Pt(14)
    r=p.add_run(text); set_run_font(r,'宋体',10.5,False)

def add_table(doc,headers,rows,widths=None,font_size=9.2):
    table=doc.add_table(rows=1,cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    hdr=table.rows[0]; set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        set_cell_text(hdr.cells[i],h,font_size,True,WD_ALIGN_PARAGRAPH.CENTER); shade_cell(hdr.cells[i],'D9E1F2')
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row): set_cell_text(cells[i],val,font_size,False,WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return table

def page_break(doc):
    p=doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

# ---------- build DOCX ----------
doc=Document(); sec=doc.sections[0]
sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(2.1); sec.right_margin=Cm(2.1); sec.header_distance=Cm(0.8); sec.footer_distance=Cm(1.0)
add_page_number(sec)
styles=doc.styles
styles['Normal'].font.name='宋体'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); styles['Normal'].font.size=Pt(12)

# Page 1
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
r=p.add_run(TITLE); set_run_font(r,'黑体',22,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4)
r=p.add_run(DATE); set_run_font(r,'宋体',12,False)
add_heading(doc,'摘    要',1)
add_para(doc,ABSTRACT,first=True,size=12,line=22)
add_para(doc,KEYWORDS,first=False,size=12,line=22)
add_heading(doc,'一、方案论证',1)
add_para(doc,'系统由小车循迹、钢球视觉测量、摆杆执行和状态显示四部分组成。选型以题面得分点、可观测性、闭环完整性和故障可控性为主，而不是单纯追求器件堆叠。',line=20)
add_table(doc,['子系统','备选方案','比较与结论'],[
['循迹检测','离散TCRT / 摄像头 / 六路红外阵列','六路I²C阵列同时提供数字状态和16位模拟量，安装、标定和加权误差计算更直接，选用LineFollower_6CH。'],
['钢球检测','电阻式/超声 / IMU推算 / K230视觉','题面要求实时画面；视觉可同时输出球心、速度和置信度，选K230作为钢球位置唯一测量。'],
['摆杆执行','RC舵机 / 直流电机 / 步进电机+编码器','舵机回差与反馈分辨率受限；步进方案便于限速与角度规划，配MS42CG闭合实际轴角。'],
['车轮驱动','开环差速 / 单循迹环 / 循迹+速度校正','在循迹PID外增加双轮速度PI，可抑制左右电机差异和电池电压变化造成的直线偏航。'],
],widths=[2.3,4.2,10.0],font_size=9.0)
add_para(doc,'最终形成“两类任务相互隔离、两套闭环分层协同”的方案：REQ-002仅运行循迹与车轮控制；REQ-003仅运行球位外环和摆杆角度内环，另一执行器组始终锁止。',line=20)
page_break(doc)

# Page 2
add_heading(doc,'二、系统理论分析与计算',1)
add_heading(doc,'1、小车循迹控制理论',2)
add_para(doc,'设左右轮线速度为vL、vR，轮距为B，则车体线速度和角速度分别为：',line=20)
add_equation(doc,'v=(vR+vL)/2，  ω=(vR−vL)/B')
add_para(doc,'六路传感器位置权值取pi∈{−1.0,−0.6,−0.2,0.2,0.6,1.0}。对白底/黑线标定值归一化得到si，循迹误差采用加权质心：',line=20)
add_equation(doc,'eL=Σ(pi·si)/Σsi−e0，  uL=Kp·eL+Ki∫eLdt+Kd·deL/dt')
add_para(doc,'控制量uL叠加到左右轮基准PWM，并在弯道降低内侧轮需求。直线段再比较10 ms窗口编码器增量；左轮A相1倍频计数乘4后与右轮硬件QEI同量纲比较，速度PI仅在软启动完成、转向量较小且两轮均有新脉冲时更新，缺脉冲或故障立即清积分。',line=20)
track_len=3+pi
add_table(doc,['验收场景','路径/时间','理论最低平均速度'],[
['A→B','1.5 m / 8 s','0.188 m/s'],
['单圈循迹停车',f'{track_len:.3f} m / 20 s','0.307 m/s'],
['动态平衡单圈',f'{track_len:.3f} m / 30 s','0.205 m/s'],
],widths=[4.5,5.5,6.5],font_size=9.5)
add_heading(doc,'2、球—杆系统控制理论',2)
add_para(doc,'将钢球视为无滑动滚动的实心球，x为相对摆杆中心的位置，θ为摆杆小角度，at为车辆沿杆方向的等效扰动。忽略高阶项时：',line=20)
add_equation(doc,'x¨≈(5/7)(g·θ−at)')
add_para(doc,'该对象近似为二重积分环节，单纯开环给角易产生越位和振荡。因此外环依据K230的球位、速度和数据置信度计算期望杆角θd，采用PD为主、弱积分消除静差，并对θd、角速度和积分限幅；内环依据MS42CG计数及连杆标定表控制D36A的STEP/DIR/EN，使实际杆角跟踪θd。',line=20)
doc.add_picture(str(FIG_CTRL),width=Cm(16.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc,'图2-1  循迹与滚球双闭环控制结构')
add_para(doc,'当视觉帧超过100 ms未更新、CRC错误、置信度不足或帧号停滞时，旧坐标立即失效；控制器清积分、摆杆回安全角并禁止继续使用预测位置。',line=19)
page_break(doc)

# Page 3
add_heading(doc,'三、电路与程序设计',1)
add_heading(doc,'1、电路与系统结构',2)
doc.add_picture(str(FIG_SYS),width=Cm(16.4)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc,'图3-1  系统总体框图')
add_table(doc,['模块','主要接口','设计作用与关键约束'],[
['LineFollower_6CH','I²C0：PA28/PA31，5 V','六路模拟/数字循迹；总线异常重试，线序和极性需实测确认。'],
['K230','UART1_RX：PA9，115200 8N1','固定14字节、CRC-8/ATM；输出球位、速度、置信度和像素坐标。'],
['左右轮/DRV8870','PB14/PB12/PA7/PB24','四路逻辑控制双H桥；电机功率线绕过扩展板，默认输出为零。'],
['左右编码器','PB4/PB5；PB10/PB11','左轮GPIO 1倍频、右轮TIMG8 QEI；用于速度估计和失速检测。'],
['D36A通道1','PA26/PA24/PB0','STEP/DIR/EN控制摆杆；必须限频、限角、限时并具硬件失能。'],
['MS42CG','PA29/PA30，PWM/Z保留','A/B角度反馈、零点与失步监测；3.3 V逻辑域。'],
['OLED','I²C1：PB3/PB2','显示模式、时间、球位和故障码；不参与闭环。'],
],widths=[3.1,4.7,8.9],font_size=8.7)
add_heading(doc,'2、电源与安全设计',2)
add_para(doc,'逻辑域、视觉模块和功率执行器分域供电并共地；车轮与步进功率不由开发板3.3 V供电。电池标称11.1 V，5 V由降压模块提供，接入前必须测量空载/带载电压与纹波。电机支路设置保险和可触及物理断能；软件上电默认PWM为0、D36A失能，控制超期、丢线、编码器停滞、过流或棕断均进入STOP。',line=20)
add_heading(doc,'3、接口完整性设计',2)
add_para(doc,'K230与MCU只采用单向位置链，视频图传链独立送至场外显示/存储设备；控制链与展示链解耦，避免显示阻塞影响实时控制。I²C0仅服务循迹模块，I²C1仅服务OLED，消除共享总线故障扩散。',line=20)
page_break(doc)

# Page 4
add_heading(doc,'三、电路与程序设计（续）',1)
add_heading(doc,'4、程序功能与分层设计',2)
add_para(doc,'软件采用“应用状态机—算法—驱动—板级安全”四层结构。应用层负责REQ-002/REQ-003互锁、计时和完成判据；算法层包含归一化循迹、PID及抗饱和；驱动层管理I²C、UART、编码器、PWM、OLED和无线诊断；板级安全层以SysTick持续强制安全输出。中断只完成采样入队和计数，解析与控制在前台周期任务中执行，降低中断抖动。',line=20)
doc.add_picture(str(FIG_FLOW),width=Cm(14.8)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc,'图3-2  主程序状态机与安全处理流程')
add_table(doc,['状态/机制','主要动作','转移或保护条件'],[
['IDLE/ARMED','输出锁止，等待KEY2或KEY3','按键去抖30 ms；只允许一个任务Owner。'],
['REQ-002运行','5 ms循迹控制，10 ms速度统计，软启动','离开A点后计圈；回到A点确认后停车并冻结时间。'],
['REQ-003运行','O→+5 cm→−5 cm分段目标','到点需同时满足位置误差、速度和稳定时间判据。'],
['FAULT/STOP','PWM=0，D36A失能，清积分并记录原因','传感器超时、通信CRC错、失线、卡滞、限位或控制超期。'],
],widths=[3.4,6.0,7.2],font_size=9.0)
add_heading(doc,'5、关键程序流程',2)
add_para(doc,'循迹数据先按白底/黑线标定归一化，若有效信号和不足则判丢线；PID微分项采用低通滤波，输出和积分均限幅。K230接收采用ISR单生产者队列，前台按55 AA帧头重同步并校验状态、序号和CRC；只有linkFresh与measurementUsable同时为真时才向球位控制器提交数据。',line=20)
add_para(doc,'停车采用“标志确认+最短离开时间+状态锁存”而非单次传感器触发，避免A点抖动重复计圈。到达COMPLETE后冻结计时并持续执行安全停车，不因显示或通信任务延迟而恢复输出。',line=20)
page_break(doc)

# Page 5
add_heading(doc,'四、测试方案与测试结果',1)
add_heading(doc,'1、测试方案与仪器',2)
add_para(doc,'测试按“静态检查→构建→逻辑域→架空低功率→静止无球→静止带球→地面低速→完整场景”逐级开展。主要仪器包括数字万用表、示波器/逻辑分析仪、直流电源及限流模块、转速/位移标尺、场外录像设备。每次记录固件哈希、接线版本、供电、电池状态、参数、原始视频和异常。',line=20)
add_heading(doc,'2、当前已取得的验证证据',2)
add_table(doc,['项目','结果','证据与结论'],[
['SysConfig 1.26.2生成','PASS','当前工程配置可生成；既有低功耗保持提示不影响编译。'],
['固件合同测试','11/11 PASS','覆盖左编码器GPIO计数、GROUP1中断入口、归一化与速度PI安全门。'],
['Debug编译/链接','PASS','TI clang 5.1.1 LTS；输出SHA-256为D68F…1E25。'],
['MotorSelfTest编译/链接','PASS','输出SHA-256为0F05…A2B6；仅证明构建，不等于硬件通过。'],
['无线诊断链','实测PASS','COM7收到“@RFTEST NF02PA LINK OK\\r\\n”共24字节，证明单向诊断链可用。'],
],widths=[4.2,3.0,9.4],font_size=8.8)
doc.add_picture(str(FIG_CAL),width=Cm(15.0)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc,'图4-1  六路循迹传感器当前标定表（代码静态值，需按赛道复标）')
add_heading(doc,'3、性能测试记录表',2)
add_table(doc,['题目要求','指标','实测结果','判定'],[
['REQ-001 图传','稳定实时显示并完整录像','待实测填写','NOT RUN'],
['REQ-002 单圈','≤20 s，停车偏差≤2 cm','时间____s；偏差____cm','待判定'],
['REQ-003 静止摆球','≤5 s；±5 cm误差≤1 cm','时间____s；误差+____/−____cm','待判定'],
['REQ-004 A→B','≤8 s；球位误差≤1 cm','时间____s；最大误差____cm','待判定'],
['REQ-005 单圈O点','≤30 s；球位误差≤1 cm','时间____s；最大误差____cm','待判定'],
['REQ-006 任意点','≤30 s；球位误差≤1 cm','目标____cm；最大误差____cm','待判定'],
],widths=[3.5,5.0,5.4,2.5],font_size=8.3)
add_para(doc,'说明：上表空白项必须由同一机械版本、同一固件与完整录像对应填写；不得用构建结果替代性能实测。',first=False,size=10.5,line=16)
page_break(doc)

# Page 6
add_heading(doc,'四、测试结果分析（续）',1)
add_heading(doc,'4、误差来源与针对性改进',2)
add_table(doc,['主要误差','影响','改进措施'],[
['左右轮参数不一致','直线漂移、停车偏差','编码器归一化速度PI；分直线/弯道参数；电池电压分档复测。'],
['传感器安装与光照','循迹误差和A点误判','固定离地高度；赛道现场复标六路白/黑值；设置有效信号和与确认时间。'],
['视觉延迟/失帧','球位超调甚至使用旧坐标','帧号、CRC、新鲜度与置信度门控；记录端到端延迟并做速度前馈。'],
['连杆回差/步进丢步','杆角模型偏差、静差增大','正反向五点标定；MS42CG闭环校正；限制角速度和加速度。'],
['供电纹波与热','复位、驱动限流或性能漂移','逻辑/功率分域、近端去耦、保险与散热；记录带载最低电压。'],
],widths=[3.2,4.5,8.7],font_size=8.8)
add_heading(doc,'五、结论',1)
add_para(doc,'本设计围绕评分最高的动态平衡任务，建立了可观测、可执行、可保护的完整闭环：六路红外阵列与差速速度校正负责稳定循迹，K230提供钢球位置，D36A与MS42CG构成摆杆角度执行内环；软件通过模式互锁、状态机、限幅、抗饱和和多源失效检测统一管理。理论上，分层双闭环能够把车体运动扰动与球位控制解耦，并以时间预算和误差预算直接对应题面指标。当前已完成资料复核、接口冻结、合同测试和编译链接验证；整车时间、停车精度和球位误差仍须按第4节表格完成实测后才能宣称达标。',line=20)
add_heading(doc,'参考文献',1)
refs=[
'[1] 全国大学生电子设计竞赛组委会. 2026年TI杯赛区赛H题：车载平衡滚球运动控制系统.',
'[2] Texas Instruments. MSPM0G3507 Mixed-Signal Microcontroller Technical Documentation.',
'[3] HiWonder/AiBlock. LineFollower_6CH V1.0 快速使用说明及原理图.',
'[4] Kendryte. K230目标检测/跟踪开发资料.',
'[5] Texas Instruments. DRV8870 Brushed DC Motor Driver Datasheet.',
'[6] D36A双路步进电机驱动器用户手册；MS42CG V2编码器用户手册.',
'[7] SSD1306 OLED控制器数据手册；nRF24L01+产品规格书.',
]
for ref in refs: add_para(doc,ref,first=False,size=10.5,line=17)
add_heading(doc,'赛前提交核验',2)
add_para(doc,'必须补齐：①六项题目性能原始数据与对应视频；②车体尺寸、摆杆长度与h≥5 cm照片；③图传接收端位于赛道外的全景证据；④最终BOM、原理图/接线图编号和固件哈希；⑤三次以上重复测试的最大误差而非仅填最好成绩。',first=False,size=10.5,line=17)

doc.save(DOCX_PATH)

# ---------- PDF mirror ----------
pdfmetrics.registerFont(TTFont('SimSun', SIMSUN))
pdfmetrics.registerFont(TTFont('SimHei', SIMHEI))
styles=getSampleStyleSheet()
body=ParagraphStyle('CNBody',fontName='SimSun',fontSize=11.3,leading=17.2,alignment=TA_LEFT,firstLineIndent=22,spaceAfter=1)
body0=ParagraphStyle('CNBody0',parent=body,firstLineIndent=0)
h1=ParagraphStyle('CNH1',fontName='SimHei',fontSize=14,leading=19,spaceBefore=2,spaceAfter=2)
h2=ParagraphStyle('CNH2',fontName='SimHei',fontSize=11.5,leading=16.5,spaceBefore=1,spaceAfter=1)
title_style=ParagraphStyle('TitleCN',fontName='SimHei',fontSize=20,leading=28,alignment=TA_CENTER,spaceAfter=5)
center=ParagraphStyle('CenterCN',fontName='SimSun',fontSize=10.5,leading=15,alignment=TA_CENTER)
equation=ParagraphStyle('EquationCN',fontName='SimSun',fontSize=11.0,leading=16,alignment=TA_CENTER,spaceBefore=1,spaceAfter=1)
small=ParagraphStyle('SmallCN',fontName='SimSun',fontSize=8.1,leading=11,alignment=TA_LEFT)
small_center=ParagraphStyle('SmallCenter',parent=small,alignment=TA_CENTER)
ref_style=ParagraphStyle('RefCN',fontName='SimSun',fontSize=8.8,leading=12,alignment=TA_LEFT,firstLineIndent=0)

page_w,page_h=A4

def footer(canvas,doc_):
    canvas.saveState(); canvas.setFont('SimSun',9); canvas.drawRightString(page_w-2.1*cm,0.9*cm,str(doc_.page)); canvas.restoreState()

def PT(text,style=body): return Paragraph(text.replace('&','&amp;').replace('\n','<br/>'),style)
def table_pdf(headers, rows, widths, fs=7.7):
    data=[[PT(h,small_center) for h in headers]]+[[PT(str(v),small_center if i==0 else small) for i,v in enumerate(row)] for row in rows]
    t=Table(data,colWidths=[w*cm for w in widths],repeatRows=1,hAlign='CENTER')
    t.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),0.45,colors.black),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#D9E1F2')),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(0,0),(-1,-1),3),('RIGHTPADDING',(0,0),(-1,-1),3),
        ('TOPPADDING',(0,0),(-1,-1),2),('BOTTOMPADDING',(0,0),(-1,-1),2)
    ])); return t

story=[]
# p1
story += [PT(TITLE,title_style),PT(DATE,center),PT('摘    要',h1),PT(ABSTRACT,body),PT(KEYWORDS,body0),PT('一、方案论证',h1),PT('系统由小车循迹、钢球视觉测量、摆杆执行和状态显示四部分组成。选型以题面得分点、可观测性、闭环完整性和故障可控性为主，而不是单纯追求器件堆叠。',body)]
story.append(table_pdf(['子系统','备选方案','比较与结论'],[
['循迹检测','离散TCRT / 摄像头 / 六路红外阵列','六路I²C阵列同时提供数字状态和16位模拟量，安装、标定和加权误差计算更直接，选用LineFollower_6CH。'],
['钢球检测','电阻式/超声 / IMU推算 / K230视觉','题面要求实时画面；视觉可同时输出球心、速度和置信度，选K230作为钢球位置唯一测量。'],
['摆杆执行','RC舵机 / 直流电机 / 步进电机+编码器','步进方案便于限速与角度规划，配MS42CG闭合实际轴角。'],
['车轮驱动','开环差速 / 单循迹环 / 循迹+速度校正','在循迹PID外增加双轮速度PI，抑制电机差异和电池变化造成的偏航。'],
],[2.2,4.0,10.1]))
story += [PT('最终形成“两类任务相互隔离、两套闭环分层协同”的方案：REQ-002仅运行循迹与车轮控制；REQ-003仅运行球位外环和摆杆角度内环，另一执行器组始终锁止。',body),PageBreak()]
# p2
story += [PT('二、系统理论分析与计算',h1),PT('1、小车循迹控制理论',h2),PT('设左右轮线速度为vL、vR，轮距为B，则车体速度关系为：',body),PT('v = (vR + vL) / 2，ω = (vR - vL) / B',equation),PT('六路传感器位置权值为{-1.0，-0.6，-0.2，0.2，0.6，1.0}。各通道按白底/黑线标定值归一化后，以加权质心计算循迹误差：',body),PT('eL = Σ(pi × si) / Σsi - e0',equation),PT('循迹控制采用PID并叠加弯道限速。直线段比较10ms编码器增量；左轮A相1倍频乘4后与右轮QEI同量纲比较。速度PI仅在软启动完成、转向量较小且两轮有新脉冲时更新，缺脉冲或故障清积分。',body)]
story.append(table_pdf(['验收场景','路径/时间','理论最低平均速度'],[['A→B','1.5 m / 8 s','0.188 m/s'],['单圈循迹停车',f'{track_len:.3f} m / 20 s','0.307 m/s'],['动态平衡单圈',f'{track_len:.3f} m / 30 s','0.205 m/s']],[4.3,5.2,6.6]))
story += [PT('2、球—杆系统控制理论',h2),PT('将钢球视为无滑动滚动的实心球，x为相对摆杆中心的位置，θ为摆杆小角度，at为车辆沿杆方向的等效扰动。忽略高阶项时：',body),PT('d²x / dt² ≈ (5/7) × (gθ - at)',equation),PT('该对象近似为二重积分环节。外环依据K230球位、速度和置信度计算期望杆角θd，采用PD为主、弱积分消除静差，并对杆角、角速度和积分限幅；内环依据MS42CG计数及连杆标定表控制D36A。',body),RLImage(str(FIG_CTRL),width=15.3*cm,height=6.6*cm),PT('图2-1  循迹与滚球双闭环控制结构',center),PT('当视觉帧超过100ms未更新、CRC错误、置信度不足或帧号停滞时，旧坐标立即失效；控制器清积分、摆杆回安全角并禁止继续使用预测位置。',body),PageBreak()]
# p3
story += [PT('三、电路与程序设计',h1),PT('1、电路与系统结构',h2),RLImage(str(FIG_SYS),width=15.5*cm,height=6.98*cm),PT('图3-1  系统总体框图',center)]
story.append(table_pdf(['模块','主要接口','设计作用与关键约束'],[
['LineFollower_6CH','I²C0：PA28/PA31，5 V','六路循迹；总线异常重试，线序和极性需实测。'],['K230','UART1_RX：PA9，115200 8N1','14字节、CRC-8/ATM；输出球位、速度和置信度。'],['DRV8870','PB14/PB12/PA7/PB24','双H桥；功率线旁路扩展板，默认零输出。'],['左右编码器','PB4/PB5；PB10/PB11','左GPIO 1倍频、右TIMG8 QEI；测速和失速检测。'],['D36A','PA26/PA24/PB0','STEP/DIR/EN；必须限频、限角、限时。'],['MS42CG','PA29/PA30','A/B角度反馈；PWM/Z保留；3.3 V逻辑域。'],['OLED','I²C1：PB3/PB2','显示模式、时间、球位和故障码。'],],[2.9,4.5,8.7],fs=7.1))
story += [PT('2、电源与安全设计',h2),PT('逻辑、视觉和功率执行器分域供电并共地；车轮与步进功率采用独立电源，不由开发板3.3V供电。电机支路设置保险和物理断能。上电默认PWM为0、D36A失能；故障均进入STOP，包括控制超期、丢线、编码器停滞、过流和棕断。',body),PT('K230控制链与场外视频展示链解耦；I2C0仅服务循迹模块，I2C1仅服务OLED，以减少共享故障。',body),PageBreak()]
# p4
story += [PT('三、电路与程序设计（续）',h1),PT('4、程序功能与分层设计',h2),PT('软件采用“应用状态机—算法—驱动—板级安全”四层结构。应用层负责任务互锁、计时和完成判据；算法层实现归一化循迹、PID和抗饱和；驱动层管理通信与执行器；安全层以SysTick持续强制安全输出。中断只完成采样入队和计数，解析与控制在前台周期任务中执行。',body),RLImage(str(FIG_FLOW),width=14.3*cm,height=7.82*cm),PT('图3-2  主程序状态机与安全处理流程',center)]
story.append(table_pdf(['状态/机制','主要动作','转移或保护条件'],[
['IDLE/ARMED','输出锁止，等待KEY2或KEY3','去抖30 ms；只允许一个任务Owner。'],['REQ-002','5 ms循迹，10 ms测速，软启动','离开A后计圈；回A确认后停车。'],['REQ-003','O→+5 cm→−5 cm分段目标','同时满足位置、速度与稳定时间。'],['FAULT/STOP','PWM=0，D36A失能，清积分','超时、CRC错、失线、卡滞或限位。'],],[3.2,5.5,7.4],fs=7.3))
story += [PT('循迹数据先按白底/黑线归一化，有效信号不足则判丢线；K230接收采用ISR队列，前台按55 AA帧头重同步并校验状态、序号和CRC，只有linkFresh与measurementUsable同时为真才提交控制。停车采用“标志确认+最短离开时间+状态锁存”，到达COMPLETE后冻结计时并持续安全停车。',body),PageBreak()]
# p5
story += [PT('四、测试方案与测试结果',h1),PT('1、测试方案与仪器',h2),PT('测试按“静态检查→构建→逻辑域→架空低功率→静止无球→静止带球→地面低速→完整场景”逐级开展。使用数字万用表、示波器/逻辑分析仪、限流电源、位移标尺和场外录像设备；每次记录固件哈希、接线、供电、参数、原始视频和异常。',body),PT('2、当前已取得的验证证据',h2)]
story.append(table_pdf(['项目','结果','证据与结论'],[
['SysConfig 1.26.2生成','PASS','当前配置可生成。'],['固件合同测试','11/11 PASS','覆盖GPIO编码器、归一化与速度PI安全门。'],['Debug构建','PASS','TI clang 5.1.1 LTS；SHA-256 D68F…1E25。'],['MotorSelfTest构建','PASS','SHA-256 0F05…A2B6；不等于硬件通过。'],['无线诊断链','实测PASS','COM7收到@RFTEST NF02PA LINK OK，共24字节。'],],[4.0,3.0,9.1],fs=7.3))
story += [Spacer(1,5),RLImage(str(FIG_CAL),width=14.8*cm,height=5.13*cm),PT('图4-1  六路循迹传感器当前标定表（代码静态值，需按赛道复标）',center),PT('3、性能测试记录表',h2)]
story.append(table_pdf(['题目要求','指标','实测结果','判定'],[
['REQ-001 图传','实时显示并完整录像','待实测填写','NOT RUN'],['REQ-002 单圈','≤20 s；≤2 cm','____s；____cm','待判定'],['REQ-003 静止摆球','≤5 s；误差≤1 cm','____s；+____/−____cm','待判定'],['REQ-004 A→B','≤8 s；误差≤1 cm','____s；____cm','待判定'],['REQ-005 单圈O点','≤30 s；误差≤1 cm','____s；____cm','待判定'],['REQ-006 任意点','≤30 s；误差≤1 cm','目标____cm；误差____cm','待判定'],],[3.3,4.5,5.2,2.8],fs=6.8))
story += [PT('空白项必须与同一机械/固件版本及完整录像对应，不得用构建结果替代性能实测。',ref_style),PageBreak()]
# p6
story += [PT('四、测试结果分析（续）',h1),PT('4、误差来源与针对性改进',h2)]
story.append(table_pdf(['主要误差','影响','改进措施'],[
['左右轮参数差异','漂移、停车偏差','速度PI；直线/弯道分参数；电池分档复测。'],['安装与光照','循迹误差、A点误判','固定高度；现场复标；设置确认时间。'],['视觉延迟/失帧','球位超调、旧坐标','帧号/CRC/新鲜度门控；记录延迟并做速度前馈。'],['连杆回差/丢步','杆角偏差、静差','正反五点标定；编码器校正；限制角速度。'],['供电纹波与热','复位、限流、漂移','分域供电、去耦、保险、散热与带载记录。'],],[3.0,4.1,8.9],fs=7.3))
story += [PT('五、结论',h1),PT('本设计围绕动态平衡高分任务，建立可观测、可执行、可保护的完整闭环：六路红外阵列与差速速度校正负责循迹，K230提供钢球位置，D36A与MS42CG构成摆杆角度执行内环；软件通过模式互锁、状态机、限幅、抗饱和和失效检测统一管理。理论上，分层双闭环能把车体运动扰动与球位控制解耦，并以时间和误差预算对应题面指标。当前已完成资料复核、接口冻结、合同测试和编译链接；整车时间、停车精度和球位误差仍须完成实测后才能宣称达标。',body),PT('参考文献',h1)]
for ref in refs: story.append(PT(ref,ref_style))
story += [PT('赛前提交核验',h2),PT('必须补齐：①六项性能原始数据与视频；②车体尺寸、摆杆长度和h≥5 cm照片；③赛道外图传接收端全景；④最终BOM、接线图编号和固件哈希；⑤三次以上重复测试的最大误差。',ref_style)]

pdf=SimpleDocTemplate(str(PDF_PATH),pagesize=A4,rightMargin=2.1*cm,leftMargin=2.1*cm,topMargin=1.55*cm,bottomMargin=1.45*cm,title=TITLE,author='2026 TI杯H题项目组')
pdf.build(story,onFirstPage=footer,onLaterPages=footer)

# basic structural evidence
import zipfile, re, hashlib
with zipfile.ZipFile(DOCX_PATH) as z:
    xml=z.read('word/document.xml')
    page_breaks=xml.count(b'w:type="page"')
    footer_parts=[n for n in z.namelist() if n.startswith('word/footer')]
print(f'DOCX={DOCX_PATH}')
print(f'PDF={PDF_PATH}')
print(f'ABSTRACT_CHARS={len(ABSTRACT)}')
print(f'PAGE_BREAKS={page_breaks}; expected=5')
print(f'FOOTERS={footer_parts}')
for p in [DOCX_PATH,PDF_PATH,FIG_SYS,FIG_CTRL,FIG_FLOW,FIG_CAL]:
    print(p.name,p.stat().st_size,hashlib.sha256(p.read_bytes()).hexdigest().upper())

